import re
import csv
from pathlib import Path


def _detect_encoding(file_path: str) -> str:
    """自动检测文件编码"""
    import charset_normalizer
    with open(file_path, "rb") as f:
        result = charset_normalizer.from_bytes(f.read())
    best = result.best()
    return best.encoding if best else "utf-8"


def extract_text_from_pdf(file_path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    texts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            texts.append(text.strip())
    return "\n\n".join(texts)


def extract_text_from_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def extract_text_from_xlsx(file_path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(file_path, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        parts.append(f"[{sheet_name}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_from_csv(file_path: str) -> str:
    encoding = _detect_encoding(file_path)
    parts = []
    with open(file_path, "r", encoding=encoding, errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            cells = [c.strip() for c in row if c.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_from_html(file_path: str) -> str:
    from bs4 import BeautifulSoup
    encoding = _detect_encoding(file_path)
    with open(file_path, "r", encoding=encoding, errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    # 移除 script 和 style 元素
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n", strip=True)
    # 合并空行
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return "\n\n".join(lines)


def extract_text_from_file(file_path: str, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in {".xlsx", ".xls"}:
        return extract_text_from_xlsx(file_path)
    elif ext == ".csv":
        return extract_text_from_csv(file_path)
    elif ext in {".html", ".htm"}:
        return extract_text_from_html(file_path)
    else:
        encoding = _detect_encoding(file_path)
        with open(file_path, "r", encoding=encoding, errors="replace") as f:
            return f.read()


def split_into_chunks(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """按段落优先的文本分块"""
    # 先按段落分割
    paragraphs = re.split(r"\n{2,}", text.strip())
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    chunks = []
    current_chunk = ""

    for para in paragraphs:
        # 如果段落本身超过 chunk_size，强制按字符切分
        if len(para) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
            # 切分超长段落
            for i in range(0, len(para), chunk_size - overlap):
                sub = para[i:i + chunk_size]
                if sub.strip():
                    chunks.append(sub.strip())
            continue

        if len(current_chunk) + len(para) + 2 <= chunk_size:
            current_chunk = (current_chunk + "\n\n" + para).strip()
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            # overlap: 将上一个 chunk 的末尾部分带入新 chunk
            overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
            current_chunk = (overlap_text + "\n\n" + para).strip()

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    return [c for c in chunks if c]
